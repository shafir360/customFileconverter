"""
docx_generator.py

This module contains business logic for creating a nicely formatted Word document
from a structured JSON object. The JSON object must follow the structure:

{
  "title": "Your Document Title",
  "sections": [
    {
      "type": "heading" | "subheading" | "paragraph" | "bullet" |
              "numbered" | "quote" | "highlight" | "bold" | "table" |
              "image" | "summary" | "questions" | "resources",
      "content": "The text content or an array of strings (for lists) or a nested array (for tables)"
    },
    ...
  ]
}

This module includes:
  - A helper function for parsing simple markdown syntax for **bold** and __highlighted__
    text.
  - A main function `build_word_document` that creates the document.
  - A function `generate_image_from_prompt` that calls an n8n webhook to generate an image.
"""

import re
import os
import uuid
import requests
from docx import Document
from docx.shared import Pt, Inches
from PIL import Image as PILImage  # Make sure Pillow is installed

def add_formatted_text(paragraph, text):
    """
    Parse the input text for simple markdown formatting markers for bold and highlight.
    
    - Bold syntax: **text**
    - Highlight syntax: __text__ (shown using italics to simulate a highlight)
    
    Splits the text based on these patterns and adds runs to the provided paragraph
    with the appropriate formatting.
    
    Parameters:
      paragraph: A docx Paragraph object where formatted text will be appended.
      text: The input string containing the markdown text.
    """
    pattern = r'(\*\*.*?\*\*|__.*?__)'
    parts = re.split(pattern, text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('__') and part.endswith('__'):
            run = paragraph.add_run(part[2:-2])
            run.italic = True   # Italic used as a visual cue for highlighting.
        else:
            paragraph.add_run(part)

def generate_image_from_prompt(prompt, webhook_url):
    """
    Calls the n8n webhook to generate an image based on the provided prompt.
    
    The webhook is expected to receive a JSON payload with the parameter "prompt".
    It should return either:
      1. Binary image data with an image content-type, or
      2. A JSON response containing an "image_url" field.
    
    The function downloads and saves the image file locally and returns the filename.
    
    Parameters:
      prompt (str): The description used to generate the image.
      webhook_url (str): The URL of the n8n webhook.
      
    Returns:
      str: The local filename of the downloaded/generated image.
      
    Raises:
      Exception: If image generation fails or if the webhook response is not as expected.
    """
    payload = {"prompt": prompt}
    response = requests.post(webhook_url, json=payload)
    if response.status_code == 200:
        content_type = response.headers.get('Content-Type', '')
        # Case 1: Webhook returns binary image data directly.
        if content_type.startswith("image/"):
            ext = content_type.split("/")[-1]
            image_filename = f"generated_image_{uuid.uuid4().hex}.{ext}"
            with open(image_filename, "wb") as f:
                f.write(response.content)
            return image_filename
        else:
            # Case 2: Webhook returns JSON with an image URL.
            try:
                data = response.json()
                image_url = data.get("image_url")
                if not image_url:
                    raise Exception("Webhook response JSON does not contain 'image_url'.")
                image_response = requests.get(image_url)
                if image_response.status_code == 200:
                    # Determine file extension from URL or default to jpg.
                    ext = image_url.split('.')[-1] if '.' in image_url else "jpg"
                    image_filename = f"generated_image_{uuid.uuid4().hex}.{ext}"
                    with open(image_filename, "wb") as f:
                        f.write(image_response.content)
                    return image_filename
                else:
                    raise Exception("Failed to download image from webhook-provided URL.")
            except Exception as e:
                raise Exception("Error processing webhook response: " + str(e))
    else:
        raise Exception("Image generation webhook error. Status code: " + str(response.status_code))

def build_word_document(json_obj, output_filename="generated_document.docx", include_images=False, report_image_errors=False, webhook_url=None):
    """
    Create a nicely formatted Word document from the provided JSON object.
    
    The document includes different formatting styles based on the section type.
    Image blocks are processed only if include_images is True and a valid webhook_url is provided.
    When processing an image block, the function calls the `generate_image_from_prompt` function to 
    retrieve the generated image, scales the image so it doesn't take up too much space (using a 
    maximum width of 6 inches), inserts the image with a caption (using the image prompt), and then 
    removes the temporary image file.
    
    If an error occurs during image generation, and if report_image_errors is True, a placeholder
    error message is added; otherwise, the image block is silently skipped.
    
    Parameters:
      json_obj (dict): The JSON object containing "title" and "sections" keys.
      output_filename (str): The name of the output Word document file.
      include_images (bool): Optional flag indicating if image sections should be included.
                             If False, image blocks will be skipped.
      report_image_errors (bool): Optional flag indicating if image generation errors should be 
                                  reported in the document. Defaults to False.
      webhook_url (str): Optional webhook URL to generate images. If not provided, image sections are skipped.
    
    Returns:
      str: The filename of the saved Word document.
    """
    doc = Document()
    # Add the document title as a top-level heading (level 0)
    doc.add_heading(json_obj.get("title", "Document"), level=0)

    image_count = 0  # <-- NEW: Initialize image counter.

    for section in json_obj.get("sections", []):
        section_type = section.get("type", "").lower()
        content = section.get("content", "")

        if section_type == "heading":
            doc.add_heading(content, level=1)
        elif section_type == "subheading":
            doc.add_heading(content, level=2)
        elif section_type == "paragraph":
            para = doc.add_paragraph()
            add_formatted_text(para, content)
        elif section_type == "bullet":
            if isinstance(content, list):
                for item in content:
                    para = doc.add_paragraph(style="List Bullet")
                    add_formatted_text(para, item)
            else:
                para = doc.add_paragraph(style="List Bullet")
                add_formatted_text(para, content)
        elif section_type == "numbered":
            if isinstance(content, list):
                for item in content:
                    para = doc.add_paragraph(style="List Number")
                    add_formatted_text(para, item)
            else:
                para = doc.add_paragraph(style="List Number")
                add_formatted_text(para, content)
        elif section_type == "quote":
            doc.add_paragraph(content, style="Intense Quote")
        elif section_type == "highlight":
            para = doc.add_paragraph(content)
        elif section_type == "bold":
            para = doc.add_paragraph()
            add_formatted_text(para, content)
        elif section_type == "table":
            if isinstance(content, list) and content and isinstance(content[0], list):
                rows = len(content)
                cols = len(content[0])
                table = doc.add_table(rows=rows, cols=cols)
                for i, row in enumerate(content):
                    cells = table.rows[i].cells
                    for j, cell_text in enumerate(row):
                        cells[j].text = str(cell_text)
            else:
                doc.add_paragraph("Table data not in expected format.")
        elif section_type == "image":
            # Enforce maximum of 3 images.
            if image_count < 3:
                if include_images and webhook_url:
                    try:
                        # Generate and download the image from the webhook.
                        image_filename = generate_image_from_prompt(content, webhook_url)
                        
                        # Use a with-statement to open the image and get dimensions.
                        max_width = 6.0  # Maximum width in inches.
                        with PILImage.open(image_filename) as im:
                            dpi = im.info.get("dpi", (96, 96))[0]  # Default DPI.
                            natural_width = im.size[0] / dpi  # Width in inches.
                            width_inches = max_width if natural_width > max_width else natural_width
                        
                        # Insert the image with the specified width.
                        doc.add_picture(image_filename, width=Inches(width_inches))
                        # Add a caption below the image using the prompt.
                        doc.add_paragraph(content, style="Caption")
                        # Remove the temporary image file.
                        os.remove(image_filename)
                        
                        image_count += 1  # <-- NEW: Increment the count after a successful image insertion.
                    except Exception as e:
                        if report_image_errors:
                            doc.add_paragraph("Image could not be generated: " + str(e))
                # If include_images is False or no webhook_url, the image block is skipped.
            else:
                # OPTIONAL: When max images reached, add a note if error reporting is enabled.
                if report_image_errors:
                    doc.add_paragraph("Image skipped: maximum number of 3 images reached.")
        elif section_type == "summary":
            para = doc.add_paragraph()
            run = para.add_run(content)
            run.bold = True
        elif section_type == "questions":
            if isinstance(content, list):
                for item in content:
                    para = doc.add_paragraph(style="List Bullet")
                    add_formatted_text(para, item)
            else:
                para = doc.add_paragraph(style="List Bullet")
                add_formatted_text(para, content)
        elif section_type == "resources":
            if isinstance(content, list):
                for item in content:
                    para = doc.add_paragraph(style="List Bullet")
                    add_formatted_text(para, item)
            else:
                para = doc.add_paragraph(style="List Bullet")
                add_formatted_text(para, content)
        else:
            para = doc.add_paragraph()
            add_formatted_text(para, content)

    # Save the document to disk.
    doc.save(output_filename)
    return output_filename

if __name__ == '__main__':
    # Example usage:
    sample_json = {
      "title": "Hospitality Business Strategy: Strategic Directions Using Porter’s and Bowman’s Frameworks",
      "sections": [
        {"type": "heading", "content": "Strategic Direction: The Foundation of Competitive Advantage"},
        {"type": "paragraph", "content": (
            "Strategic direction refers to the actionable plans an organization adopts to achieve "
            "its long-term goals and vision. It aligns resources, operations, and people toward common "
            "objectives, ensuring the organization thrives in a competitive environment. After analyzing "
            "internal strengths, external opportunities, and industry competition, leaders must choose a "
            "strategic direction that builds a **sustainable competitive advantage**—a unique position "
            "that competitors cannot easily replicate."
         )},
        {"type": "summary", "content": (
            "Strategic direction is crucial for aligning an organization's resources and operations "
            "towards long-term goals, ensuring a sustainable competitive advantage in a competitive environment."
         )},
        {"type": "heading", "content": "Porter’s Generic Strategies"},
        {"type": "paragraph", "content": "Michael Porter’s framework identifies three primary strategies for gaining competitive advantage:"},
        {"type": "numbered", "content": [
            "**Cost Leadership**: Becoming the lowest-cost provider in the industry while maintaining acceptable quality.",
            "**Differentiation**: Offering unique products/services that justify premium pricing.",
            "**Focus Strategy**: Targets a niche market segment through either cost focus or differentiation focus."
        ]},
        {"type": "image", "content": (
            "A visual representation of Porter’s Generic Strategies showing Cost Leadership, Differentiation, "
            "and Focus Strategy as distinct paths to competitive advantage in the hospitality industry."
        )},
        {"type": "summary", "content": (
            "Porter’s Generic Strategies provide three paths to competitive advantage: Cost Leadership, "
            "Differentiation, and Focus Strategy, each with distinct applications and implications in the hospitality industry."
        )},
        {"type": "heading", "content": "Bowman’s Strategy Clock: A Complementary Framework"},
        {"type": "paragraph", "content": (
            "Bowman’s model expands on Porter by analyzing competitive strategies based on **price** and "
            "**perceived value**. The eight positions help organizations avoid risky strategies and align with market needs."
        )},
        {"type": "bullet", "content": [
            "**Low Price/Low Value**: Risky; may lead to price wars.",
            "**Low Price**: Similar to cost leadership.",
            "**Hybrid**: Balances moderate price and value.",
            "**Differentiation**: Higher value without excessive pricing.",
            "**Focused Differentiation**: Premium pricing for exceptional value.",
            "**Increased Price/Standard Value**: Risky; customers reject overpriced average offerings.",
            "**Increased Price/Low Value**: Guaranteed failure.",
            "**Standard/Commodity**: No competitive edge."
        ]},
        {"type": "summary", "content": (
            "Bowman’s Strategy Clock offers a nuanced view of competitive strategies by linking price to perceived value, "
            "helping organizations identify viable and risky positions in the market."
        )},
        {"type": "heading", "content": "Key Takeaways"},
        {"type": "bullet", "content": [
            "**Porter** simplifies strategy into three paths, while **Bowman** provides granularity by linking price to value.",
            "A hotel chain might use cost leadership for budget brands and differentiation for luxury lines.",
            "Strategic pitfalls include over-reliance on cost-cutting eroding quality and excessive differentiation alienating price-sensitive customers."
        ]},
        {"type": "summary", "content": (
            "Evaluating Porter’s and Bowman’s models enables hospitality organizations to choose a strategic direction "
            "that leverages their strengths and meets market demands for lasting advantage."
        )},
        {"type": "questions", "content": [
            "How does cost leadership differ from differentiation in Porter’s Generic Strategies?",
            "What are the risks associated with the Increased Price/Standard Value position in Bowman’s Strategy Clock?",
            "Can you provide an example of a hospitality business successfully implementing a Focus Strategy?"
        ]}
      ]
    }
    # To test image generation, you can supply a valid webhook URL via the 'webhook_url' parameter.
    # Here, report_image_errors is set to True for demonstration.
    output_file = build_word_document(
        sample_json,
        output_filename="sample_output.docx",
        include_images=True,
        report_image_errors=True,
        webhook_url="https://your-valid-webhook-url.example.com/path"  # Replace with actual URL when testing
    )
    print(f"Document '{output_file}' created successfully.")
